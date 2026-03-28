"""Document ingestion functionality for document datasets."""

import os
import hashlib
from typing import Any, Dict, List, Optional

import structlog
from qdrant_client import QdrantClient
from sentence_transformers import SentenceTransformer

from .chunking import DocumentChunker

logger = structlog.get_logger()


class DocumentIngester:
    """Ingests documents for search and retrieval."""

    def __init__(
        self,
        qdrant_client: QdrantClient,
        embedder: SentenceTransformer,
        chunker: DocumentChunker,
        collection_name: str = "document_resources",
    ):
        """Initialize document ingester.
        
        Args:
            qdrant_client: Qdrant client instance
            embedder: Sentence transformer instance
            chunker: Document chunker instance
            collection_name: Qdrant collection name
        """
        self.qdrant_client = qdrant_client
        self.embedder = embedder
        self.chunker = chunker
        self.collection_name = collection_name

    async def initialize_collection(self) -> None:
        """Initialize Qdrant collection for document resources."""
        try:
            # Check if collection exists
            collections = self.qdrant_client.get_collections()
            collection_names = [c.name for c in collections.collections]
            
            if self.collection_name not in collection_names:
                # Create collection
                self.qdrant_client.create_collection(
                    collection_name=self.collection_name,
                    vectors_config={
                        "size": self.embedder.get_sentence_embedding_dimension(),
                        "distance": "Cosine",
                    },
                )
                logger.info("Created Qdrant collection", collection=self.collection_name)
            else:
                logger.info("Qdrant collection already exists", collection=self.collection_name)
                
        except Exception as e:
            logger.error("Failed to initialize collection", error=str(e))
            raise

    async def ingest_document(
        self,
        filename: str,
        content: bytes,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """Ingest a single document.
        
        Args:
            filename: Document filename
            content: Document content
            metadata: Optional document metadata
            
        Returns:
            Ingestion results
        """
        try:
            # Parse document
            parsed_doc = self._parse_document(filename, content)
            
            # Chunk document
            chunks = self.chunker.chunk_document(parsed_doc)
            
            # Process chunks
            chunks_created = 0
            for chunk in chunks:
                # Create resource ID
                resource_id = self._create_resource_id(
                    filename=filename,
                    chunk_index=chunk["chunk_index"],
                    page_number=chunk.get("page_number", 0),
                )
                
                # Create metadata
                chunk_metadata = {
                    "filename": filename,
                    "file_type": self._get_file_type(filename),
                    "page_number": chunk.get("page_number", 0),
                    "chunk_index": chunk["chunk_index"],
                    "chunk_size": len(chunk["content"]),
                    **chunk.get("metadata", {}),
                    **(metadata or {}),
                }
                
                # Create provenance
                provenance = {
                    "source": filename,
                    "ingested_at": self._get_timestamp(),
                    "content_hash": self._calculate_content_hash(chunk["content"]),
                    "file_size": len(content),
                }
                
                # Generate embedding
                embedding = self.embedder.encode(chunk["content"]).tolist()
                
                # Store in Qdrant
                self.qdrant_client.upsert(
                    collection_name=self.collection_name,
                    points=[
                        {
                            "id": resource_id,
                            "vector": embedding,
                            "payload": {
                                "content": chunk["content"],
                                "metadata": chunk_metadata,
                                "provenance": provenance,
                            },
                        }
                    ],
                )
                
                chunks_created += 1
            
            results = {
                "filename": filename,
                "chunks_created": chunks_created,
                "file_size": len(content),
                "content_hash": self._calculate_content_hash(content.decode('utf-8', errors='ignore')),
            }
            
            logger.info("Document ingested successfully", **results)
            return results
            
        except Exception as e:
            logger.error("Document ingestion failed", filename=filename, error=str(e))
            raise

    async def ingest_documents_batch(
        self,
        files: List[Any],
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """Ingest multiple documents.
        
        Args:
            files: List of document files
            metadata: Optional metadata for all documents
            
        Returns:
            Batch ingestion results
        """
        try:
            files_processed = 0
            total_chunks_created = 0
            
            for file in files:
                # Read file content
                content = await file.read()
                
                # Ingest document
                result = await self.ingest_document(
                    filename=file.filename,
                    content=content,
                    metadata=metadata,
                )
                
                files_processed += 1
                total_chunks_created += result["chunks_created"]
            
            results = {
                "files_processed": files_processed,
                "chunks_created": total_chunks_created,
            }
            
            logger.info("Batch document ingestion completed", **results)
            return results
            
        except Exception as e:
            logger.error("Batch document ingestion failed", error=str(e))
            raise

    def _parse_document(self, filename: str, content: bytes) -> Dict[str, Any]:
        """Parse document content.
        
        Args:
            filename: Document filename
            content: Document content
            
        Returns:
            Parsed document
        """
        file_type = self._get_file_type(filename)
        
        if file_type == "pdf":
            return self._parse_pdf(content)
        elif file_type == "docx":
            return self._parse_docx(content)
        elif file_type == "txt":
            return self._parse_txt(content)
        elif file_type == "html":
            return self._parse_html(content)
        else:
            # Fallback to text parsing
            return self._parse_txt(content)

    def _parse_pdf(self, content: bytes) -> Dict[str, Any]:
        """Parse PDF document.
        
        Args:
            content: PDF content
            
        Returns:
            Parsed document
        """
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(stream=content, filetype="pdf")
            
            pages = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                
                if text.strip():
                    pages.append({
                        "page_number": page_num + 1,
                        "content": text,
                        "metadata": {
                            "page_width": page.rect.width,
                            "page_height": page.rect.height,
                        },
                    })
            
            doc.close()
            
            return {
                "type": "pdf",
                "pages": pages,
                "total_pages": len(pages),
            }
            
        except ImportError:
            logger.warning("PyMuPDF not available, falling back to text parsing")
            return self._parse_txt(content)
        except Exception as e:
            logger.error("PDF parsing failed", error=str(e))
            return self._parse_txt(content)

    def _parse_docx(self, content: bytes) -> Dict[str, Any]:
        """Parse DOCX document.
        
        Args:
            content: DOCX content
            
        Returns:
            Parsed document
        """
        try:
            from docx import Document
            import io
            
            doc = Document(io.BytesIO(content))
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            content_text = "\n".join(paragraphs)
            
            return {
                "type": "docx",
                "content": content_text,
                "metadata": {
                    "paragraph_count": len(paragraphs),
                },
            }
            
        except ImportError:
            logger.warning("python-docx not available, falling back to text parsing")
            return self._parse_txt(content)
        except Exception as e:
            logger.error("DOCX parsing failed", error=str(e))
            return self._parse_txt(content)

    def _parse_txt(self, content: bytes) -> Dict[str, Any]:
        """Parse text document.
        
        Args:
            content: Text content
            
        Returns:
            Parsed document
        """
        try:
            text = content.decode('utf-8', errors='ignore')
            
            return {
                "type": "txt",
                "content": text,
                "metadata": {
                    "character_count": len(text),
                    "line_count": len(text.splitlines()),
                },
            }
            
        except Exception as e:
            logger.error("Text parsing failed", error=str(e))
            return {
                "type": "txt",
                "content": "",
                "metadata": {},
            }

    def _parse_html(self, content: bytes) -> Dict[str, Any]:
        """Parse HTML document.
        
        Args:
            content: HTML content
            
        Returns:
            Parsed document
        """
        try:
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return {
                "type": "html",
                "content": text,
                "metadata": {
                    "title": soup.title.string if soup.title else None,
                    "character_count": len(text),
                },
            }
            
        except ImportError:
            logger.warning("BeautifulSoup not available, falling back to text parsing")
            return self._parse_txt(content)
        except Exception as e:
            logger.error("HTML parsing failed", error=str(e))
            return self._parse_txt(content)

    def _get_file_type(self, filename: str) -> str:
        """Get file type from filename.
        
        Args:
            filename: Filename
            
        Returns:
            File type
        """
        _, ext = os.path.splitext(filename.lower())
        
        type_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'docx',
            '.txt': 'txt',
            '.html': 'html',
            '.htm': 'html',
        }
        
        return type_map.get(ext, 'txt')

    def _create_resource_id(
        self,
        filename: str,
        chunk_index: int,
        page_number: int = 0,
    ) -> str:
        """Create unique resource ID.
        
        Args:
            filename: Document filename
            chunk_index: Chunk index
            page_number: Page number
            
        Returns:
            Unique resource ID
        """
        id_string = f"{filename}:{chunk_index}:{page_number}"
        return hashlib.sha256(id_string.encode()).hexdigest()

    def _calculate_content_hash(self, content: str) -> str:
        """Calculate content hash.
        
        Args:
            content: Content to hash
            
        Returns:
            Content hash
        """
        return hashlib.sha256(content.encode('utf-8')).hexdigest()

    def _get_timestamp(self) -> str:
        """Get current timestamp."""
        from datetime import datetime
        return datetime.utcnow().isoformat()











